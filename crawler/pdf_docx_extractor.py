"""Trích xuất văn bản thuần từ các tệp PDF / DOCX / DOC.

Thử nhiều backend (theo thứ tự):
  * PDF  -> ``pdfplumber`` (tốt nhất cho tiếng Việt), dự phòng bằng ``PyPDF2``, rồi ``pymupdf``.
  * DOCX -> ``python-docx``.
  * DOC  -> ``antiword`` (tệp nhị phân hệ thống) nếu có, nếu không thì bỏ qua.

Hàm trả về một *danh sách trang* (mỗi trang là một chuỗi). Với DOCX, toàn bộ
tài liệu được trả về dưới dạng một "trang" vì DOCX không có khái niệm trang
thực sự. Bên gọi không quan tâm đến ranh giới trang có thể chỉ cần dùng
``"\n\n".join(pages)``.
"""

from __future__ import annotations

import logging
import os
import subprocess
from pathlib import Path
from typing import List, Optional

logger = logging.getLogger(__name__)


def _extract_pdf(path: str) -> List[str]:
    pages: List[str] = []
    # 1) pdfplumber
    try:
        import pdfplumber  # type: ignore

        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                txt = page.extract_text() or ""
                pages.append(txt)
        if any(p.strip() for p in pages):
            return pages
    except Exception as exc:  # pragma: no cover - cố gắng tối đa
        logger.debug("pdfplumber failed on %s: %s", path, exc)

    # 2) PyPDF2
    try:
        from PyPDF2 import PdfReader  # type: ignore

        reader = PdfReader(path)
        for page in reader.pages:
            pages.append(page.extract_text() or "")
        if any(p.strip() for p in pages):
            return pages
    except Exception as exc:  # pragma: no cover
        logger.debug("PyPDF2 failed on %s: %s", path, exc)

    # 3) pymupdf (fitz)
    try:
        import fitz  # type: ignore

        doc = fitz.open(path)
        for page in doc:
            pages.append(page.get_text())
        doc.close()
        if any(p.strip() for p in pages):
            return pages
    except Exception as exc:  # pragma: no cover
        logger.debug("pymupdf failed on %s: %s", path, exc)

    logger.warning("All PDF back-ends failed for %s", path)
    return pages or [""]


def _extract_docx(path: str) -> List[str]:
    try:
        from docx import Document  # type: ignore

        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        # Đồng thời lấy văn bản từ bảng (các quy định thường dùng bảng).
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    txt = cell.text.strip()
                    if txt:
                        paragraphs.append(txt)
        return ["\n".join(paragraphs)]
    except Exception as exc:  # pragma: no cover
        logger.warning("python-docx failed on %s: %s", path, exc)
        return [""]


def _extract_doc(path: str) -> List[str]:
    """Trích xuất văn bản từ tệp .doc cũ bằng antiword nếu có."""
    try:
        result = subprocess.run(
            ["antiword", path], capture_output=True, timeout=60, check=False
        )
        if result.returncode == 0:
            txt = result.stdout.decode("utf-8", errors="ignore")
            if txt.strip():
                return [txt]
    except FileNotFoundError:
        logger.warning("antiword not installed - skipping .doc file: %s", path)
    except Exception as exc:  # pragma: no cover
        logger.warning("antiword failed on %s: %s", path, exc)
    return [""]


def extract_text_from_file(path: str) -> List[str]:
    """Trả về danh sách chuỗi văn bản theo trang của tài liệu đã cho.

    Trả về ``[""]`` (một trang trống) khi trích xuất thất bại hoàn toàn để bên
    gọi có thể lặp mà không cần xử lý trường hợp đặc biệt.
    """
    p = Path(path)
    if not p.exists():
        logger.warning("File does not exist: %s", path)
        return [""]

    ext = p.suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(path)
    if ext == ".docx":
        return _extract_docx(path)
    if ext == ".doc":
        return _extract_doc(path)
    logger.warning("Unsupported file extension: %s", path)
    return [""]


def extract_text(path: str) -> str:
    """Hàm bao tiện dụng trả về một chuỗi đã nối duy nhất."""
    return "\n\n".join(extract_text_from_file(path))
